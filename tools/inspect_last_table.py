from pathlib import Path

from docx import Document


def main() -> None:
    doc_path = Path("data/_debug_nodup.docx") if Path("data/_debug_nodup.docx").exists() else Path("data/output_330974803.docx")
    doc = Document(doc_path)
    print("tables:", len(doc.tables))
    if not doc.tables:
        return
    t = doc.tables[-1]
    print("last table rows:", len(t.rows), "cols:", len(t.columns))
    for i, row in enumerate(t.rows):
        hit_any = False
        for j, cell in enumerate(row.cells):
            txt = cell.text
            cgen = txt.count("Генеральный директор")
            cdir = txt.count("Директор")
            if cgen or cdir:
                hit_any = True
                print(f"ROW {i} CELL {j}: count_gen={cgen} count_dir={cdir}")
        if hit_any:
            print("---")


if __name__ == "__main__":
    main()


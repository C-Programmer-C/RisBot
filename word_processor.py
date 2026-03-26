import re
import logging
from pathlib import Path
from typing import Any
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from decimal import Decimal

logger = logging.getLogger(__name__)


def format_director_string(director_fio: str, is_general_director: bool, max_length: int = 47) -> str:
    """
    Формирует строку директора в формате: "Генеральный директор________________ / Иванов И.И."
    или "Директор________________ / Иванов И.И."
    
    Подчеркивания добавляются между должностью и разделителем " / " для заполнения до max_length.
    Если строка не помещается, подчеркивания убираются или ФИО сокращается.
    
    Args:
        director_fio: ФИО директора
        is_general_director: Является ли генеральным директором
        max_length: Максимальная длина строки (по умолчанию 42)
        
    Returns:
        str: Отформатированная строка директора
    """
    if not director_fio or director_fio == "ВЫ НЕ УКАЗАЛИ ДИРЕКТОРА":
        director_fio = "________________"
    
    # Если в ФИО/названии есть "ИП" — роль "директор" не пишем
    has_ip = "ИП" in str(director_fio).upper()

    # Определяем должность
    director_title = "" if has_ip else ("Генеральный директор" if is_general_director else "Директор")
    
    # Длина фиксированных частей: должность + " / " + ФИО
    title_length = len(director_title)
    fio_length = len(director_fio)
    separator = " / "
    separator_length = len(separator)
    
    # Вычисляем сколько места осталось для подчеркиваний
    total_fixed = title_length + separator_length + fio_length
    available_for_underscores = max_length - total_fixed
    
    if available_for_underscores > 0:
        # Есть место для подчеркиваний - добавляем их между должностью и разделителем
        underscores = "_" * available_for_underscores
        return f"{director_title}{underscores}{separator}{director_fio}"
    elif available_for_underscores == 0:
        # Точно помещается без подчеркиваний
        return f"{director_title}{separator}{director_fio}"
    else:
        # Не помещается - нужно сократить
        # Сначала пробуем без подчеркиваний
        if total_fixed <= max_length:
            return f"{director_title}{separator}{director_fio}"
        
        # Если все равно не помещается, сокращаем ФИО
        available_for_fio = max_length - title_length - separator_length
        if available_for_fio > 0:
            director_fio = director_fio[:available_for_fio]
            return f"{director_title}{separator}{director_fio}"
        else:
            # Если даже должность не помещается, возвращаем только должность (обрезаем)
            return director_title[:max_length]


async def get_director_data(lead_task_id: int) -> tuple[str, bool]:
    """
    Получает данные директора из задачи лида в Pyrus.
    Ищет поля с id 53 (ФИО директора) и id 54 (Генеральный директор).
    
    Args:
        lead_task_id: ID задачи лида
        
    Returns:
        tuple[str, bool]: (ФИО директора, является ли генеральным директором)
    """
    try:
        from pyrus_api_service import api_request
        
        # Получаем данные задачи из Pyrus API
        task_data = await api_request(
            method="GET",
            endpoint=f"/tasks/{lead_task_id}"
        )
        
        if not isinstance(task_data, dict):
            logger.warning(f"Invalid task data format for task {lead_task_id}")
            return "ВЫ НЕ УКАЗАЛИ ДИРЕКТОРА", False
        
        task = task_data.get("task", {})
        fields = task.get("fields", [])
        
        if not fields:
            logger.warning(f"No fields found in task {lead_task_id}")
            return "ВЫ НЕ УКАЗАЛИ ДИРЕКТОРА", False
        
        # Ищем поле с id 53 (ФИО директора)
        director_fio = "ВЫ НЕ УКАЗАЛИ ДИРЕКТОРА"
        is_general_director = False
        
        for field in fields:
            field_id = field.get("id")
            
            # Поле с id 53 - ФИО директора
            if field_id == 53:
                value = extract_field_value(field)
                if value and str(value).strip():
                    director_fio = str(value).strip()
                    logger.info(f"Found director FIO (id=53): {director_fio}")
            
            # Поле с id 54 - Генеральный директор (checked/unchecked)
            elif field_id == 54:
                field_value = field.get("value")
                field_type = field.get("type", "")
                
                # Проверяем, является ли значение "checked"
                if isinstance(field_value, str):
                    is_general_director = field_value.lower() in ["checked", "true", "1", "yes"]
                elif isinstance(field_value, dict):
                    # Может быть структура с checked/unchecked
                    checked_value = field_value.get("value", field_value.get("checked", field_value.get("choice_id")))
                    if isinstance(checked_value, bool):
                        is_general_director = checked_value
                    elif isinstance(checked_value, str):
                        is_general_director = checked_value.lower() in ["checked", "true", "1", "yes"]
                    elif isinstance(checked_value, int):
                        # Может быть choice_id (0 = unchecked, >0 = checked)
                        is_general_director = checked_value > 0
                elif isinstance(field_value, bool):
                    is_general_director = field_value
                elif isinstance(field_value, int):
                    # Может быть choice_id (0 = unchecked, >0 = checked)
                    is_general_director = field_value > 0
                elif field_value is None or field_value == "":
                    # Пустое значение = unchecked
                    is_general_director = False
                
                logger.info(f"Found general director field (id=54, type={field_type}): value={field_value}, is_general={is_general_director}")
        
        logger.info(f"Director data for task {lead_task_id}: FIO='{director_fio}', is_general={is_general_director}")
        return director_fio, is_general_director
        
    except Exception as e:
        logger.error(f"Error getting director data for task {lead_task_id}: {e}", exc_info=True)
        return "ВЫ НЕ УКАЗАЛИ ДИРЕКТОРА", False


def format_money(value: float | str | int, with_spaces: bool = False) -> str:
    """Форматирует денежное значение с 2 знаками после запятой (русский формат)"""
    try:
        if isinstance(value, str):
            value = float(value.replace(" ", "").replace(",", "."))
        num_value = float(value)
        
        integer_part = int(num_value)
        decimal_part = abs(num_value - integer_part)
        
        if with_spaces:
            integer_str = f"{integer_part:,}".replace(",", " ")
        else:
            integer_str = str(integer_part)
        
        decimal_str = f"{decimal_part:.2f}".split(".")[1]
        
        return f"{integer_str},{decimal_str}"
    except (ValueError, TypeError):
        return str(value)


def extract_field_value(field: dict[str, Any], return_numeric: bool = False) -> str | float:
    """Извлекает значение из поля в зависимости от его типа"""
    value = field.get("value")
    
    if value is None:
        return 0.0 if return_numeric else ""
    
    field_type = field.get("type", "")
    
    if field_type == "catalog":
        if isinstance(value, dict):
            values_list = value.get("values", [])
            if values_list:
                return str(values_list[0]) if values_list[0] else ""
            rows = value.get("rows", [])
            if rows and rows[0]:
                return str(rows[0][0]) if rows[0][0] else ""
        return str(value)
    
    if field_type == "multiple_choice":
        if isinstance(value, dict):
            choice_names = value.get("choice_names", [])
            return ", ".join(choice_names) if choice_names else ""
        return str(value)
    
    if field_type == "person":
        if isinstance(value, dict):
            first_name = value.get("first_name", "")
            last_name = value.get("last_name", "")
            return f"{first_name} {last_name}".strip()
        return str(value)
    
    if field_type == "date":
        return str(value)
    
    if field_type == "money" or field_type == "number":
        if return_numeric:
            try:
                return float(value) if value else 0.0
            except (ValueError, TypeError):
                return 0.0
        return str(value)
    
    if isinstance(value, (list, dict)):
        return str(value)
    
    return str(value)


def get_full_address(short_name: str) -> str:
    """Преобразует короткое название адреса в полный адрес"""
    if not short_name:
        return ""
    
    address_mapping = {
        "Северская": "Краснодарский край, Северский район, станица Северская, ул. Запорожская, д. 79",
        "Ленинский": "Краснодарский край, Абинский район, Ольгинское сельское поселение, хутор Ленинский,  улица Ленина, 2/2",
        "Староджерелиевская": "Краснодарский край, Красноармейский м.р-н, Староджерелиевское с.п., ст-ца, ул. Октябрьская, зем.уч. 68А, (координаты: 45.475575 38.308048)",
        "Стерлитамак": "Республика Башкортостан г. Стерлитамак ул. Глинки 1"
    }
    
    short_name_clean = short_name.strip()
    return address_mapping.get(short_name_clean, short_name)


def format_delivery_type(value: str) -> str:
    """Форматирует тип доставки для отображения"""
    if not value:
        return ""
    
    value_lower = value.lower()
    if "силами поставщика" in value_lower or "поставщика" in value_lower:
        return "Транспорт поставщика"
    elif "самовывоз" in value_lower:
        return "Самовывоз"
    
    return value


def number_to_words_russian(value: float | int) -> str:
    """Преобразует число в прописью на русском языке (для денег)"""
    ones = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
    teens = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
             'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
    
    def convert_three_digits(num: int) -> str:
        """Конвертирует трехзначное число"""
        result = []
        if num >= 100:
            result.append(hundreds[num // 100])
            num %= 100
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
        elif num >= 10:
            result.append(teens[num - 10])
            return ' '.join(result)
        if num > 0:
            result.append(ones[num])
        return ' '.join(result).strip()
    
    # Обрабатываем рубли
    rubles = int(value)
    kopecks = int(round((value - rubles) * 100))
    
    parts = []
    
    # Миллионы
    if rubles >= 1000000:
        millions = rubles // 1000000
        parts.append(convert_three_digits(millions))
        if millions == 1:
            parts.append('миллион')
        elif millions in [2, 3, 4]:
            parts.append('миллиона')
        else:
            parts.append('миллионов')
        rubles %= 1000000
    
    # Тысячи
    if rubles >= 1000:
        thousands = rubles // 1000
        parts.append(convert_three_digits(thousands))
        if thousands == 1:
            parts.append('тысяча')
        elif thousands in [2, 3, 4]:
            parts.append('тысячи')
        else:
            parts.append('тысяч')
        rubles %= 1000
    
    # Сотни, десятки, единицы
    if rubles > 0:
        parts.append(convert_three_digits(rubles))
    
    rubles_text = ' '.join(parts).strip()
    
    # Определяем правильную форму слова "рубль"
    rubles_last_digit = int(value) % 10
    rubles_last_two = int(value) % 100
    if rubles_last_two in [11, 12, 13, 14]:
        rubles_word = 'рублей'
    elif rubles_last_digit == 1:
        rubles_word = 'рубль'
    elif rubles_last_digit in [2, 3, 4]:
        rubles_word = 'рубля'
    else:
        rubles_word = 'рублей'
    
    # Копейки
    kopecks_word = 'копеек'
    if kopecks == 1:
        kopecks_word = 'копейка'
    elif kopecks in [2, 3, 4]:
        kopecks_word = 'копейки'
    
    kopecks_text = f"{kopecks:02d} {kopecks_word}"
    
    return f"{rubles_text} {rubles_word} {kopecks_text}"


def number_to_words_russian_days(value: int) -> str:
    """Преобразует число дней в прописью на русском языке (для дней)"""
    ones = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять']
    teens = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
             'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
    
    def convert_three_digits(num: int) -> str:
        """Конвертирует трехзначное число"""
        result = []
        if num >= 100:
            result.append(hundreds[num // 100])
            num %= 100
        if num >= 20:
            result.append(tens[num // 10])
            num %= 10
        elif num >= 10:
            result.append(teens[num - 10])
            return ' '.join(result)
        if num > 0:
            result.append(ones[num])
        return ' '.join(result).strip()
    
    if value == 0:
        return 'ноль'
    
    parts = []
    
    if value >= 1000:
        thousands = value // 1000
        parts.append(convert_three_digits(thousands))
        if thousands == 1:
            parts.append('тысяча')
        elif thousands in [2, 3, 4]:
            parts.append('тысячи')
        else:
            parts.append('тысяч')
        value %= 1000
    
    if value > 0:
        parts.append(convert_three_digits(value))
    
    return ' '.join(parts).strip()


def get_payment_conditions(payment_type: str, payment_date: str, deferral_amount: str, system_days: str) -> str:
    """Формирует условия оплаты на основе типа оплаты
    
    Args:
        payment_type: Тип оплаты
        payment_date: Дата условия оплаты (из поля id 132)
        deferral_amount: Сумма отсрочки
        system_days: Количество дней отсрочки
    """
    if not payment_type:
        return ""
    
    payment_type_lower = payment_type.lower()
    
    # Форматируем дату условия оплаты
    date_formatted = "ДАТА НЕ УКАЗАНА"
    if payment_date:
        try:
            date_formatted = format_date_russian(payment_date)
        except:
            date_formatted = payment_date if payment_date else "ДАТА НЕ УКАЗАНА"
    
    if "предоплата" in payment_type_lower:
        # Для предоплаты форматируем дату с "г." вместо "года"
        date_formatted_prepayment = date_formatted.replace(" года", " г.")
        return f"по 100 (сто) % предоплате, путем перечисления денежных средств на расчетный счет Поставщика до {date_formatted_prepayment}."
    
    if "постоплат" in payment_type_lower or "отсрочк" in payment_type_lower:
        # Если дни не указаны, используем альтернативный текст с датой из поля id 132
        if not system_days or str(system_days).strip() == "":
            # Форматируем дату условия оплаты (из поля id 132)
            date_formatted = "ДАТА НЕ УКАЗАНА"
            if payment_date:
                try:
                    date_formatted = format_date_russian(payment_date)
                    # Убираем "года" в конце и добавляем "г."
                    date_formatted = date_formatted.replace(" года", " г.")
                except:
                    date_formatted = payment_date if payment_date else "ДАТА НЕ УКАЗАНА"
            
            return f"по 100 (сто) % предоплате, путем перечисления денежных средств на расчетный счет Поставщика до {date_formatted}."
        
        # Если дни указаны, используем стандартный текст
        # Получаем дни из поля "Количество календарных дней со дня поставки товара на склад Покупателя"
        days_text = system_days if system_days else "7"
        try:
            days_int = int(float(system_days)) if system_days else 7
            days_words = number_to_words_russian_days(days_int)
        except:
            days_words = "семь" if not system_days else str(system_days)
            days_text = "7" if not system_days else str(system_days)
        
        return f"Покупатель осуществляет оплату за каждую партию товара в течении {days_text} ({days_words}) календарных дней со дня поставки товара на склад Покупателя. Поставка товара на склад Покупателя подтверждается датой подписания покупателем товаросопроводительных накладных документов (ТТН/УПД)."
    
    return ""


def format_loading_method(method: str) -> str:
    """Форматирует способ погрузки для текста"""
    if not method:
        return ""
    
    method_lower = method.lower()
    if "паллеты" in method_lower or "паллет" in method_lower:
        return "на паллетах"
    elif "навалом" in method_lower:
        return "навалом"
    elif "россыпь" in method_lower or "россыпью" in method_lower:
        return "россыпью"
    
    return method


def get_product_description(products_list: list[dict[str, str]]) -> str:
    """Формирует описание товаров в соответствии с наименованием"""
    base_text = "Товар отгружается в соответствии с наименованием товара."
    
    if not products_list:
        return base_text
    
    product_descriptions = []
    for product_info in products_list:
        product_name = product_info.get("name", "").strip()
        total_kg = product_info.get("kg", "")
        
        if product_name:
            if total_kg:
                total_kg_clean = str(total_kg).strip()
                product_descriptions.append(f"{product_name} {total_kg_clean} кг.")
            else:
                product_descriptions.append(f"{product_name}.")
    
    if product_descriptions:
        return f"{base_text} {' '.join(product_descriptions)}"
    
    return base_text


def get_loading_description(loading_method: str, packaging_kg: str) -> str:
    """Формирует описание способа отгрузки товара для одного товара"""
    if not loading_method and not packaging_kg:
        return ""
    
    loading_method_lower = loading_method.lower() if loading_method else ""
    loading_text = format_loading_method(loading_method)
    packaging_text = packaging_kg if packaging_kg else ""
    
    is_scattered = "россыпь" in loading_method_lower or "россыпью" in loading_method_lower
    is_bulk = "навалом" in loading_method_lower
    
    if is_scattered:
        return f"россыпью"
    elif is_bulk and packaging_kg:
        return f"в мешках по {packaging_text} кг навалом"
    elif loading_method and packaging_kg:
        return f"в п/п мешках по {packaging_text} кг {loading_text}"
    elif packaging_kg:
        return f"в п/п мешках по {packaging_text} кг"
    elif loading_method:
        return f"{loading_text}"
    
    return ""


def get_loading_description_multiple(products_loading: list[dict[str, str]]) -> str:
    """Формирует описание способа отгрузки для нескольких товаров"""
    if not products_loading:
        return ""
    
    descriptions = []
    for product_info in products_loading:
        loading_method = product_info.get("loading_method", "")
        packaging_kg = product_info.get("packaging_kg", "")
        desc = get_loading_description(loading_method, packaging_kg)
        if desc:
            descriptions.append(desc)
    
    if not descriptions:
        return ""
    
    # Объединяем описания
    if len(descriptions) == 1:
        return f"Товар отгружается {descriptions[0]}."
    
    # Если несколько - объединяем через " и "
    result = "Товар отгружается "
    if len(descriptions) == 2:
        result += f"{descriptions[0]} и {descriptions[1]}."
    else:
        # Для трех и более: "X, Y и Z"
        result += ", ".join(descriptions[:-1]) + f" и {descriptions[-1]}."
    
    return result


def format_date_russian(date_str: str) -> str:
    """Форматирует дату из формата YYYY-MM-DD или DD.MM.YYYY в формат «DD» месяца YYYY года."""
    if not date_str:
        return date_str
    
    try:
        # Пробуем формат DD.MM.YYYY
        if '.' in date_str:
            date_obj = datetime.strptime(date_str, "%d.%m.%Y")
        else:
            # Пробуем формат YYYY-MM-DD
            date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        
        months = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]
        day = date_obj.day
        month = months[date_obj.month - 1]
        year = date_obj.year
        return f"«{day}» {month} {year} года"
    except (ValueError, TypeError):
        return date_str


def format_number_with_spaces(value: int | float | str) -> str:
    """Форматирует число с пробелами для тысяч"""
    try:
        if isinstance(value, str):
            value = int(float(value))
        num_value = int(float(value))
        return f"{num_value:,}".replace(",", " ")
    except (ValueError, TypeError):
        return str(value)


def generate_shipping_date_text(table_fields: list[dict[str, Any]]) -> str:
    """Генерирует текст для поля 'Дата отгрузки' на основе анализа строк таблицы"""
    rows_data = []
    
    for table_field in table_fields:
        table_value = table_field.get("value", [])
        if isinstance(table_value, list) and table_value:
            for row_data in table_value:
                cells_data = row_data.get("cells", [])
                if not cells_data:
                    continue
                
                row_info = {}
                for cell_data in cells_data:
                    name = cell_data.get("name", "")
                    if name:
                        value = extract_field_value(cell_data)
                        if name == "Прайс":
                            row_info["product"] = str(value).strip()
                        elif name == "Дата отгрузки":
                            row_info["date"] = str(value).strip()
                        elif name == "всего кг":
                            try:
                                row_info["kg"] = float(value) if value else 0.0
                            except (ValueError, TypeError):
                                row_info["kg"] = 0.0
                
                # Добавляем строку, если есть дата (продукт может отсутствовать)
                if row_info.get("date"):
                    # Если продукт отсутствует, используем пустую строку
                    if "product" not in row_info:
                        row_info["product"] = ""
                    rows_data.append(row_info)
    
    if not rows_data:
        return ""
    
    unique_products = set(row["product"] for row in rows_data)
    unique_dates = set(row["date"] for row in rows_data)
    
    # Закомментирована старая логика с "включительно" - теперь это обрабатывается через поле id: 128
    # if len(unique_products) > 1 and len(unique_dates) == 1:
    #     date_str = list(unique_dates)[0]
    #     formatted_date = format_date_russian(date_str)
    #     return f"{formatted_date}, включительно"
    
    # Возвращаем только дату, без количества
    unique_dates_list = sorted(list(unique_dates))
    if len(unique_dates_list) == 1:
        formatted_date = format_date_russian(unique_dates_list[0])
        return formatted_date
    
    date_lines = []
    for date_str in unique_dates_list:
        formatted_date = format_date_russian(date_str)
        date_lines.append(formatted_date)
    
    return "\n".join(date_lines)


def create_fields_map(fields: list[dict[str, Any]]) -> dict[str, str]:
    """Создает словарь для быстрого поиска значений по имени поля"""
    fields_map = {}

    payment_type_value = ""
    payment_date_value = ""  # Дата условия оплаты (id: 132)
    deferral_amount_value = ""
    system_days_value = ""
    
    # Переменные для обработки "Дата отгрузки спецификации" (id: 122) и галочки "Включительно" (id: 128)
    shipping_spec_date_value = ""  # Дата отгрузки спецификации (id: 122)
    shipping_spec_inclusive = False  # Включительно (Дата отгрузки) (id: 128)

    def process_field(field: dict[str, Any], is_nested: bool = False):
        """Обрабатывает одно поле (рекурсивно для вложенных полей)"""
        name = field.get("name")
        field_id = field.get("id")
        if not name:
            return

        value = extract_field_value(field)

        # Извлекаем [SYSTEM] DATE LOADING для других целей
        if name == "[SYSTEM] DATE LOADING":
            # Добавляем в fields_map только если значение не пустое
            if value and str(value).strip():
                # Форматируем дату в нужный формат
                date_formatted = format_date_russian(str(value).strip())
                fields_map[name] = date_formatted
                fields_map[name.strip()] = date_formatted
            return

        # Обрабатываем "Дата отгрузки спецификации" (id: 122)
        if field_id == 122 or (name == "Дата отгрузки" and is_nested):
            nonlocal shipping_spec_date_value
            shipping_spec_date_value = value
            return  # Не добавляем в fields_map здесь, обработаем позже

        # Обрабатываем галочку "Включительно (Дата отгрузки)" (id: 128)
        if field_id == 128:
            nonlocal shipping_spec_inclusive
            if isinstance(value, str):
                shipping_spec_inclusive = value.lower() == "checked"
            elif isinstance(value, bool):
                shipping_spec_inclusive = value
            elif isinstance(value, int):
                shipping_spec_inclusive = value > 0
            return  # Не добавляем в fields_map здесь

        # Пропускаем "Дата отгрузки" из верхнего уровня - она формируется из таблицы
        if name == "Дата отгрузки" and not is_nested:
            return

        # Пропускаем поле id 130 "Товар отгружается в соответствии с наименованием товара"
        # Оно обрабатывается отдельно в process_word_template через check_manual_loading_text
        if field_id == 130 or (name == "Товар отгружается в соответствии с наименованием товара" and is_nested):
            return

        # Собираем поля для условий оплаты
        if name == "Тип оплаты":
            nonlocal payment_type_value
            payment_type_value = value
        elif field_id == 132 or name == "Дата условия оплаты":
            # Поле id 132 - Дата условия оплаты
            nonlocal payment_date_value
            payment_date_value = value
        elif name == "Сумма отсрочки":
            nonlocal deferral_amount_value
            deferral_amount_value = value
        elif (
            name
            == "Количество календарных дней со дня поставки товара на склад Покупателя"
        ):
            nonlocal system_days_value
            system_days_value = value

        # Удалена старая логика с get_full_address для "Адрес отгрузки"
        # Теперь адрес берется из таблицы id 109, поле "Адрес доставки" id 110
        if name == "Тип доставки":
            value = format_delivery_type(value)
        elif name == "Организация":
            org_value = str(value)
            org_value = org_value.replace('\\"', '"')
            parts = org_value.split('"')
            if len(parts) > 1:
                result_parts = [parts[0]]
                for i in range(1, len(parts)):
                    if i % 2 == 1:
                        result_parts.append('«' + parts[i])
                    else:
                        result_parts.append('»' + parts[i])
                org_value = ''.join(result_parts)
            value = org_value

        # Добавляем точку в конце для "Комментарий по заявке", если ее нет
        if name == "Комментарий по заявке" and value:
            value_str = str(value).strip()
            if value_str and not value_str.endswith(('.', '!', '?')):
                value = value_str + "."
            else:
                value = value_str

        fields_map[name] = value
        fields_map[name.strip()] = value

    # Обрабатываем все поля верхнего уровня
    for field in fields:
        field_type = field.get("type", "")
        field_name = field.get("name", "")

        # Обрабатываем обычные поля
        process_field(field, is_nested=False)

        # Обрабатываем вложенные поля в полях типа "title"
        if field_type == "title":
            field_value = field.get("value")
            if isinstance(field_value, dict):
                nested_fields = field_value.get("fields", [])
                if nested_fields:
                    for nested_field in nested_fields:
                        process_field(nested_field, is_nested=True)

    # Проверяем, установлена ли галочка "Условия оплаты (Текст вручную)" (id: 137)
    # Если установлена, используем текст из поля id 138 вместо автоматически сформированного
    use_manual_payment_conditions = False
    manual_payment_conditions_text = ""
    field_137 = find_field_by_id(fields, 137)
    if field_137:
        field_value = field_137.get("value")
        # Проверяем, установлена ли галочка
        if isinstance(field_value, str):
            use_manual_payment_conditions = field_value.lower() == "checked"
        elif isinstance(field_value, bool):
            use_manual_payment_conditions = field_value
        elif isinstance(field_value, int):
            use_manual_payment_conditions = field_value > 0

    # Если галочка установлена, ищем поле id 138
    if use_manual_payment_conditions:
        field_138 = find_field_by_id(fields, 138)
        if field_138:
            manual_payment_conditions_text = extract_field_value(field_138)
            if manual_payment_conditions_text:
                manual_payment_conditions_text = str(manual_payment_conditions_text).strip()

    # Проверяем, установлена ли галочка "Включительно (Условия оплаты)" (id: 142)
    payment_inclusive = False
    field_142 = find_field_by_id(fields, 142)
    if field_142:
        field_value = field_142.get("value")
        # Проверяем, установлена ли галочка
        if isinstance(field_value, str):
            payment_inclusive = field_value.lower() == "checked"
        elif isinstance(field_value, bool):
            payment_inclusive = field_value
        elif isinstance(field_value, int):
            payment_inclusive = field_value > 0

    # Формируем условия оплаты
    if use_manual_payment_conditions and manual_payment_conditions_text:
        # Используем ручной текст из поля id 138
        payment_conditions_text = manual_payment_conditions_text
        # Если галочка "Включительно" установлена и тип оплаты - предоплата, добавляем "включительно."
        if payment_inclusive and payment_type_value and "предоплата" in payment_type_value.lower():
            # Убираем точку в конце, если есть, и добавляем "включительно."
            payment_conditions_text = payment_conditions_text.rstrip('.') + ", включительно."
        # Убираем лишние точки в конце (если две или более точек, оставляем одну)
        while payment_conditions_text.endswith('..'):
            payment_conditions_text = payment_conditions_text[:-1]
        fields_map["Условия оплаты"] = payment_conditions_text
        fields_map["Условия оплаты".strip()] = payment_conditions_text
        logger.info(f"Using manual payment conditions text from field id 138: {payment_conditions_text}")
    else:
        # Используем автоматически сформированный текст
        payment_conditions = get_payment_conditions(payment_type_value, payment_date_value, deferral_amount_value, system_days_value)
        if payment_conditions:
            # Если галочка "Включительно" установлена и тип оплаты - предоплата, добавляем "включительно."
            if payment_inclusive and payment_type_value and "предоплата" in payment_type_value.lower():
                # Убираем точку в конце, если есть, и добавляем "включительно."
                payment_conditions = payment_conditions.rstrip('.') + ", включительно."
            # Убираем лишние точки в конце (если две или более точек, оставляем одну)
            while payment_conditions.endswith('..'):
                payment_conditions = payment_conditions[:-1]
            fields_map["Условия оплаты"] = payment_conditions
            fields_map["Условия оплаты".strip()] = payment_conditions

    # Обрабатываем "Дата отгрузки спецификации" (id: 122) с проверкой галочки "Включительно" (id: 128)
    if shipping_spec_date_value:
        try:
            date_formatted = format_date_russian(str(shipping_spec_date_value).strip())
            # Добавляем точку в конце
            date_formatted = date_formatted.rstrip('.') + "."
            # Если галочка "Включительно" установлена, добавляем "включительно" в конце
            if shipping_spec_inclusive:
                date_formatted = date_formatted.rstrip('.') + " включительно."
            fields_map["Дата отгрузки спецификации"] = date_formatted
            fields_map["Дата отгрузки спецификации".strip()] = date_formatted
        except Exception as e:
            logger.error(f"Error formatting shipping spec date: {e}", exc_info=True)
            fields_map["Дата отгрузки спецификации"] = str(shipping_spec_date_value)
            fields_map["Дата отгрузки спецификации".strip()] = str(shipping_spec_date_value)

    # Проверяем, установлена ли галочка "Включить адрес выгрузки" (id: 136)
    show_shipping_address = False
    field_136 = find_field_by_id(fields, 136)
    if field_136:
        field_value = field_136.get("value")
        # Проверяем, установлена ли галочка
        if isinstance(field_value, str):
            show_shipping_address = field_value.lower() == "checked"
        elif isinstance(field_value, bool):
            show_shipping_address = field_value
        elif isinstance(field_value, int):
            show_shipping_address = field_value > 0

    # Обрабатываем "Адрес отгрузки" из таблицы id 109, поле "Адрес доставки" id 110
    # Только если галочка id 136 установлена
    shipping_address = ""
    if show_shipping_address:
        for field in fields:
            field_id = field.get("id")
            field_type = field.get("type", "")
            if field_id == 109 and field_type == "table":
                table_value = field.get("value", [])
                if isinstance(table_value, list) and table_value:
                    # Берем первую строку таблицы
                    for row_data in table_value:
                        cells_data = row_data.get("cells", [])
                        if cells_data:
                            for cell_data in cells_data:
                                cell_id = cell_data.get("id")
                                cell_name = cell_data.get("name", "")
                                # Ищем поле "Адрес доставки" с id 110
                                if cell_id == 110 or cell_name == "Адрес доставки":
                                    address_value = extract_field_value(cell_data)
                                    if address_value:
                                        shipping_address = str(address_value).strip()
                                        break
                            if shipping_address:
                                break
                        if shipping_address:
                            break
                if shipping_address:
                    break

    if shipping_address:
        # Изменяем ключ с "Адрес отгрузки" на "Адрес (Организация)"
        fields_map["Адрес (Организация)"] = shipping_address
        fields_map["Адрес (Организация)".strip()] = shipping_address
        logger.info(f"Using shipping address from table id 109: {shipping_address}")
    elif not show_shipping_address:
        # Если галочка не установлена, не добавляем поле вообще (или устанавливаем пустое значение)
        # Пустое значение приведет к удалению параграфа в replace_placeholders_in_paragraph
        fields_map["Адрес (Организация)"] = ""
        fields_map["Адрес (Организация)".strip()] = ""
        logger.info("Shipping address field disabled (checkbox id 136 is unchecked)")

    # Обрабатываем "Адрес отгрузки спецификации" (id: 133) и связанные поля
    try:
        shipping_spec_address_text = build_shipping_spec_address_text(fields)
        if shipping_spec_address_text:
            fields_map["Адрес отгрузки спецификации"] = shipping_spec_address_text
            fields_map["Адрес отгрузки спецификации".strip()] = shipping_spec_address_text
            logger.info(f"Using shipping spec address text: {shipping_spec_address_text}")
    except Exception as e:
        logger.error(f"Error processing shipping spec address: {e}", exc_info=True)
    
    # Обрабатываем поля для ЖД доставки (id: 139, 140, 141)
    # Эти поля отображаются только если "Тип доставки" (id: 7) = "Силами поставщика ЖД"
    field_7 = find_field_by_id(fields, 7)
    is_railway_delivery = False
    if field_7:
        field_value = field_7.get("value")
        # Для multiple_choice значение хранится в choice_names
        if isinstance(field_value, dict) and "choice_names" in field_value:
            choice_names = field_value.get("choice_names", [])
            if choice_names:
                delivery_type = str(choice_names[0]).strip()
                if "Силами поставщика ЖД" in delivery_type or "ЖД" in delivery_type:
                    is_railway_delivery = True
        elif isinstance(field_value, str):
            if "Силами поставщика ЖД" in field_value or "ЖД" in field_value:
                is_railway_delivery = True
    
    # Если тип доставки = "Силами поставщика ЖД", добавляем поля id 139, 140, 141
    if is_railway_delivery:
        # Грузоотправитель (id: 139)
        field_139 = find_field_by_id(fields, 139)
        if field_139:
            value_139 = extract_field_value(field_139)
            if value_139 and str(value_139).strip():
                fields_map["Грузоотправитель"] = str(value_139).strip()
                fields_map["Грузоотправитель".strip()] = str(value_139).strip()
            else:
                # Если поле пустое, устанавливаем пустую строку (параграф будет удален)
                fields_map["Грузоотправитель"] = ""
                fields_map["Грузоотправитель".strip()] = ""
        else:
            fields_map["Грузоотправитель"] = ""
            fields_map["Грузоотправитель".strip()] = ""
        
        # Станция назначения (id: 140)
        field_140 = find_field_by_id(fields, 140)
        if field_140:
            value_140 = extract_field_value(field_140)
            if value_140 and str(value_140).strip():
                fields_map["Станция назначения"] = str(value_140).strip()
                fields_map["Станция назначения".strip()] = str(value_140).strip()
            else:
                # Если поле пустое, устанавливаем пустую строку (параграф будет удален)
                fields_map["Станция назначения"] = ""
                fields_map["Станция назначения".strip()] = ""
        else:
            fields_map["Станция назначения"] = ""
            fields_map["Станция назначения".strip()] = ""
        
        # Грузополучатель (id: 141)
        field_141 = find_field_by_id(fields, 141)
        if field_141:
            value_141 = extract_field_value(field_141)
            if value_141 and str(value_141).strip():
                fields_map["Грузополучатель"] = str(value_141).strip()
                fields_map["Грузополучатель".strip()] = str(value_141).strip()
            else:
                # Если поле пустое, устанавливаем пустую строку (параграф будет удален)
                fields_map["Грузополучатель"] = ""
                fields_map["Грузополучатель".strip()] = ""
        else:
            fields_map["Грузополучатель"] = ""
            fields_map["Грузополучатель".strip()] = ""
    else:
        # Если тип доставки не "Силами поставщика ЖД", устанавливаем пустые значения
        # чтобы параграфы с этими плейсхолдерами были удалены
        fields_map["Грузоотправитель"] = ""
        fields_map["Грузоотправитель".strip()] = ""
        fields_map["Станция назначения"] = ""
        fields_map["Станция назначения".strip()] = ""
        fields_map["Грузополучатель"] = ""
        fields_map["Грузополучатель".strip()] = ""

    return fields_map


def find_table_fields(fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Находит все поля типа table"""
    return [field for field in fields if field.get("type") == "table"]


def replace_placeholders_in_paragraph(paragraph, fields_map: dict[str, str], font_size: int | None = None, bold: bool | None = None):
    """Заменяет плейсхолдеры в параграфе. Возвращает True если параграф нужно удалить"""
    text = paragraph.text
    
    if "${" not in text:
        return False
    
    import re
    # Используем finditer для замены каждого плейсхолдера только один раз
    pattern = r'\$\{([^}]+)\}'
    matches = list(re.finditer(pattern, text))
    
    # Заменяем с конца, чтобы индексы не сдвигались
    for match in reversed(matches):
        placeholder_name = match.group(1)
        placeholder_name_clean = placeholder_name.strip()
        placeholder_full = match.group(0)  # Полный плейсхолдер с ${}
        
        value = None
        if placeholder_name_clean in fields_map:
            value = fields_map[placeholder_name_clean]
        elif placeholder_name in fields_map:
            value = fields_map[placeholder_name]
        
        if value is not None:
            value_str = str(value).strip()
            if not value_str:
                return True
            # Заменяем только это конкретное вхождение
            start, end = match.span()
            text = text[:start] + value_str + text[end:]
        else:
            # Удаляем плейсхолдер, если значение не найдено
            start, end = match.span()
            text = text[:start] + text[end:]
    
    if text != paragraph.text:
        paragraph.clear()
        run = paragraph.add_run(text)
        if font_size is not None:
            run.font.size = Pt(font_size)
        if bold is not None:
            run.font.bold = bold
    
    return False


def replace_placeholders_in_table_cell(cell, fields_map: dict[str, str]):
    """Заменяет плейсхолдеры в ячейке таблицы"""
    for paragraph in cell.paragraphs:
        replace_placeholders_in_paragraph(paragraph, fields_map, font_size=10, bold=False)


def create_table_row_map(cells_data: list[dict[str, Any]]) -> dict[str, Any]:
    """Создает словарь для быстрого поиска значений ячеек таблицы по их name"""
    row_map = {}
    for cell_data in cells_data:
        name = cell_data.get("name", "")
        if name:
            value = extract_field_value(cell_data)
            # Удалена старая логика с get_full_address для "Адрес отгрузки"
            # Теперь адрес берется из таблицы id 109, поле "Адрес доставки" id 110
            row_map[name] = value
            row_map[name.strip()] = value
            row_map[f"{name}_raw"] = cell_data
    return row_map


def group_rows_by_product(rows_data: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Группирует строки по наименованию товара (Прайс) и суммирует стоимость"""
    grouped = {}
    
    for row_map in rows_data:
        product_name = row_map.get("Прайс", "").strip()
        if not product_name:
            continue
        
        if product_name not in grouped:
            grouped[product_name] = row_map.copy()
            grouped[product_name]["_count"] = 1
            try:
                price_value = extract_field_value(grouped[product_name].get("Цена_raw", {}), return_numeric=True)
                grouped[product_name]["_total_price"] = float(price_value) if price_value else 0.0

                kg_value = extract_field_value(grouped[product_name].get("всего кг_raw", {}), return_numeric=True)
                grouped[product_name]["_total_kg"] = float(kg_value) if kg_value else 0.0
                
                price_per_kg_value = extract_field_value(grouped[product_name].get("Цена за кг (Дост)_raw", {}), return_numeric=True)
                grouped[product_name]["_price_per_kg"] = float(price_per_kg_value) if price_per_kg_value else 0.0
            except:
                grouped[product_name]["_total_price"] = 0.0
                grouped[product_name]["_total_kg"] = 0.0
                grouped[product_name]["_price_per_kg"] = 0.0
        else:
            grouped[product_name]["_count"] += 1
            try:
                price_value = extract_field_value(row_map.get("Цена_raw", {}), return_numeric=True)
                grouped[product_name]["_total_price"] += float(price_value) if price_value else 0.0
            except:
                pass
            try:
                kg_value = extract_field_value(row_map.get("всего кг_raw", {}), return_numeric=True)
                grouped[product_name]["_total_kg"] += float(kg_value) if kg_value else 0.0
            except:
                pass
    
    result = []
    for product_name, row_data in grouped.items():
        total_kg = row_data.get("_total_kg", 0.0)
        if isinstance(total_kg, (int, float)) and float(total_kg).is_integer():
            row_data["всего кг"] = str(int(float(total_kg)))
        else:
            row_data["всего кг"] = str(total_kg)
        row_data["Цена"] = format_money(row_data["_total_price"], with_spaces=True)
        row_data["Цена за кг (Дост)"] = format_money(row_data.get("_price_per_kg", 0.0), with_spaces=False)
        result.append(row_data)
    
    return result


def find_template_row(table) -> tuple[int, list] | None:
    """Находит строку-шаблон с плейсхолдерами в таблице"""
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if "${" in cell.text:
                return (row_idx, row)
    return None


def replace_placeholders_in_row(row, row_map: dict[str, Any], row_number: int | None = None):
    """Заменяет плейсхолдеры в строке таблицы значениями из row_map"""
    # Определяем, какие значения нужно установить
    values_to_set = {}
    
    # Номер п/п
    if row_number is not None:
        values_to_set[0] = f"{row_number}."
    
    # Остальные значения из row_map
    for name, value in row_map.items():
        if name.startswith("_") or name.endswith("_raw"):
            continue
        
        # Определяем, в какой столбец это значение идет
        # Это зависит от структуры таблицы в шаблоне
        # Прайс - обычно второй столбец (индекс 1)
        # всего кг - обычно третий столбец (индекс 2)
        # Цена за кг (Дост) - обычно четвертый столбец (индекс 3)
        # Цена - обычно пятый столбец (индекс 4)
        if name == "Прайс":
            values_to_set[1] = str(value)
        elif name == "всего кг":
            values_to_set[2] = str(value)
        elif name == "Цена за кг (Дост)":
            values_to_set[3] = str(value)
        elif name == "Цена":
            values_to_set[4] = str(value)
    
    # Устанавливаем значения напрямую
    for cell_idx, cell in enumerate(row.cells):
        cell.paragraphs.clear()
        if cell_idx in values_to_set:
            cell.add_paragraph(values_to_set[cell_idx])


def find_total_row(table) -> tuple[int, list, int] | None:
    """Находит строку с 'Итого' и индекс последней ячейки (где сумма)"""
    for row_idx, row in enumerate(table.rows):
        has_total = False
        for cell in row.cells:
            if "Итого" in cell.text:
                has_total = True
                break
        
        if has_total:
            last_cell_idx = len(row.cells) - 1
            return (row_idx, row, last_cell_idx)
    return None


def insert_table_data(doc: Document, table_name: str, table_data: list[dict[str, Any]]):
    """Вставляет данные таблицы из JSON в Word документ"""
    logger.debug(f"insert_table_data called for table '{table_name}' with {len(table_data)} rows")
    if not table_data:
        logger.debug("No table data to insert")
        return
    
    for table_idx, table in enumerate(doc.tables):
        logger.debug(f"Checking table {table_idx}, has {len(table.rows)} rows")
        template_info = find_template_row(table)
        if template_info is None:
            logger.debug(f"Table {table_idx} has no template row with placeholders")
            continue
        
        template_row_idx, template_row = template_info
        logger.debug(f"Found template row at index {template_row_idx}")
        
        # Сохраняем оригинальное содержимое ячеек шаблона (до замены плейсхолдеров)
        # Важно: сохраняем ДО любых изменений
        template_cell_texts = []
        for cell in template_row.cells:
            # Получаем чистый текст ячейки
            cell_text = ""
            for paragraph in cell.paragraphs:
                cell_text += paragraph.text
            template_cell_texts.append(cell_text.strip())
        
        logger.debug(f"Saved template cell texts: {template_cell_texts}")
        
        # Сохраняем оригинальную строку-шаблон XML для копирования структуры
        # Важно: делаем это ДО любых изменений в template_row
        from copy import deepcopy
        original_template_tr = deepcopy(template_row._tr)
        
        total_row_info = find_total_row(table)
        total_row_idx = total_row_info[0] if total_row_info else None
        total_cell_idx = total_row_info[2] if total_row_info else None
        
        rows_to_insert = []
        for row_data in table_data:
            cells_data = row_data.get("cells", [])
            if not cells_data:
                continue
            
            row_map = create_table_row_map(cells_data)
            rows_to_insert.append(row_map)
        
        if not rows_to_insert:
            continue
        
        grouped_rows = group_rows_by_product(rows_to_insert)
        
        if not grouped_rows:
            continue
        
        total_sum = 0.0
        
        for i, row_map in enumerate(grouped_rows):
            row_number = i + 1
            
            if i == 0:
                # Заменяем плейсхолдеры в первой строке (шаблонной)
                # Используем тот же подход, что и для остальных строк
                # Полностью очищаем все ячейки
                for cell in template_row.cells:
                    # Используем оба способа очистки для надежности
                    cell.text = ""
                    cell.paragraphs.clear()
                
                # Устанавливаем значения напрямую
                # Номер п/п
                if len(template_row.cells) > 0:
                    template_row.cells[0].add_paragraph(f"{row_number}.")
                
                # Прайс (наименование товара)
                product_name = row_map.get("Прайс", "")
                if len(template_row.cells) > 1 and product_name:
                    template_row.cells[1].add_paragraph(str(product_name))
                
                # всего кг (количество)
                total_kg = row_map.get("всего кг", "")
                if len(template_row.cells) > 2 and total_kg:
                    template_row.cells[2].add_paragraph(str(total_kg))
                
                # Цена за кг (Дост)
                price_per_kg = row_map.get("Цена за кг (Дост)", "")
                if len(template_row.cells) > 3 and price_per_kg:
                    template_row.cells[3].add_paragraph(str(price_per_kg))
                
                # Цена (стоимость товара)
                price = row_map.get("Цена", "")
                if len(template_row.cells) > 4 and price:
                    template_row.cells[4].add_paragraph(str(price))
            else:
                # Добавляем новую строку, копируя структуру из первой строки (шаблонной)
                # но мы полностью заменим содержимое
                new_row = table.add_row()
                
                # Полностью очищаем все ячейки новой строки
                for cell in new_row.cells:
                    # Используем оба способа очистки для надежности
                    cell.text = ""
                    cell.paragraphs.clear()
                
                # Устанавливаем значения напрямую из row_map
                # Номер п/п
                if len(new_row.cells) > 0:
                    new_row.cells[0].add_paragraph(f"{row_number}.")
                
                # Прайс (наименование товара)
                product_name = row_map.get("Прайс", "")
                if len(new_row.cells) > 1 and product_name:
                    new_row.cells[1].add_paragraph(str(product_name))
                
                # всего кг (количество)
                total_kg = row_map.get("всего кг", "")
                if len(new_row.cells) > 2 and total_kg:
                    new_row.cells[2].add_paragraph(str(total_kg))
                
                # Цена за кг (Дост)
                price_per_kg = row_map.get("Цена за кг (Дост)", "")
                if len(new_row.cells) > 3 and price_per_kg:
                    new_row.cells[3].add_paragraph(str(price_per_kg))
                
                # Цена (стоимость товара)
                price = row_map.get("Цена", "")
                if len(new_row.cells) > 4 and price:
                    new_row.cells[4].add_paragraph(str(price))
                
                logger.debug(f"Processed row {row_number}: product='{product_name}', price='{price}'")
            
            try:
                total_price = row_map.get("_total_price", 0.0)
                total_sum += float(total_price) if total_price else 0.0
            except (ValueError, TypeError):
                pass
        
        # Перемещаем строку "Итого" в конец таблицы (после всех товаров)
        if total_row_info and total_row_idx is not None and total_cell_idx is not None:
            # Находим строку "Итого"
            total_row_idx_final = None
            total_row_obj = None
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    if "Итого" in cell.text:
                        total_row_idx_final = row_idx
                        total_row_obj = row
                        break
                if total_row_idx_final is not None:
                    break
            
            if total_row_idx_final is not None and total_row_obj is not None:
                # Если "Итого" не в конце, перемещаем её
                if total_row_idx_final < len(table.rows) - 1:
                    # Удаляем строку "Итого" из текущей позиции
                    table._tbl.remove(total_row_obj._tr)
                    
                    # Добавляем строку "Итого" в конец
                    new_total_row = table.add_row()
                    total_row_obj = new_total_row
                
                # Заполняем строку "Итого" правильно: первая ячейка (объединенная) - "Итого", последняя - сумма
                formatted_total = format_money(total_sum, with_spaces=True)
                
                # Объединяем все средние ячейки с первой ячейкой
                # Это создаст 2 ячейки: первая (объединенная) с "Итого", последняя с суммой
                num_cells = len(total_row_obj.cells)
                logger.debug(f"Total row has {num_cells} cells before merging, need to reduce to 2")
                
                if num_cells > 2:
                    try:
                        # Объединяем все ячейки кроме последней с первой
                        # Нужно объединить num_cells - 2 ячеек со первой
                        merge_attempts = 0
                        max_merges = num_cells - 2  # Нужно объединить num_cells - 2 раз
                        while len(total_row_obj.cells) > 2 and merge_attempts < max_merges:
                            # Всегда объединяем первую со второй (после каждого merge индексы сдвигаются)
                            total_row_obj.cells[0].merge(total_row_obj.cells[1])
                            merge_attempts += 1
                            logger.debug(f"Merge attempt {merge_attempts}: now have {len(total_row_obj.cells)} cells")
                        
                        logger.debug(f"Finished merging: now have {len(total_row_obj.cells)} cells after {merge_attempts} merges")
                    except Exception as e:
                        logger.error(f"Error merging cells in 'Итого' row: {e}", exc_info=True)
                        # Продолжаем без объединения, просто очистим средние ячейки
                        for j in range(1, len(total_row_obj.cells) - 1):
                            total_row_obj.cells[j].text = ""
                
                # Убеждаемся, что у нас есть хотя бы 2 ячейки
                if len(total_row_obj.cells) < 2:
                    logger.warning(f"Total row has only {len(total_row_obj.cells)} cells, expected at least 2")
                
                # Сначала устанавливаем текст в ячейки в правильном порядке
                # Первая ячейка (объединенная) - "Итого"
                total_row_obj.cells[0].text = "Итого"
                logger.debug(f"Set first cell to 'Итого', now has {len(total_row_obj.cells)} cells")
                
                # Последняя ячейка - сумма (используем индекс -1)
                if len(total_row_obj.cells) > 1:
                    total_cell = total_row_obj.cells[-1]
                    total_cell.text = ""  # Очищаем ячейку
                    if total_cell.paragraphs:
                        total_cell.paragraphs[0].clear()
                        total_cell.paragraphs[0].add_run(formatted_total)
                    else:
                        total_cell.add_paragraph().add_run(formatted_total)
                    logger.debug(f"Set total cell (index {len(total_row_obj.cells)-1}) value to: {formatted_total}")
                else:
                    logger.warning("Total row has only one cell, appending sum to 'Итого'")
                    # Если только одна ячейка, добавляем сумму к "Итого"
                    total_row_obj.cells[0].text = f"Итого {formatted_total}"
        
        return


def extract_table_fields_to_map(table_fields: list[dict[str, Any]], all_fields: list[dict[str, Any]] | None = None) -> dict[str, str]:
    """Извлекает значения полей из таблиц и добавляет их в fields_map"""
    table_fields_map = {}
    
    products_dict = {}  # Словарь для группировки товаров по названию (суммируем кг)
    products_loading = []  # Список для сбора способов погрузки и упаковки всех товаров
    
    # Проверяем, установлена ли галочка для "В соответствии с наименованием" (id: 131 с именем "Текст вручную")
    # Если установлена, не формируем автоматический текст для "В соответствии с наименованием"
    use_manual_product_description = False
    if all_fields:
        is_manual, _ = check_manual_loading_text(all_fields)
        use_manual_product_description = is_manual
    
    for table_field in table_fields:
        table_value = table_field.get("value", [])
        if isinstance(table_value, list) and table_value:
            for row_data in table_value:
                cells_data = row_data.get("cells", [])
                if not cells_data:
                    continue
                
                product_name = ""
                total_kg = ""
                loading_method_value = ""
                packaging_kg_value = ""
                
                for cell_data in cells_data:
                    name = cell_data.get("name", "")
                    if name:
                        value = extract_field_value(cell_data)
                        # Удалена старая логика с get_full_address для "Адрес отгрузки"
                        # Теперь адрес берется из таблицы id 109, поле "Адрес доставки" id 110
                        if name == "Тип доставки":
                            value = format_delivery_type(value)
                        elif name == "Способ погрузки":
                            loading_method_value = value
                        elif name == "Упаковка кг.":
                            packaging_kg_value = value
                        elif name == "Прайс":
                            product_name = str(value).strip()
                        elif name == "всего кг":
                            try:
                                total_kg = float(value) if value else 0.0
                            except (ValueError, TypeError):
                                total_kg = 0.0
                        
                        table_fields_map[name] = value
                        table_fields_map[name.strip()] = value
                
                # Группируем товары по названию, суммируя количество
                if product_name:
                    if product_name in products_dict:
                        products_dict[product_name] += total_kg
                    else:
                        products_dict[product_name] = total_kg
                    
                    # Сохраняем информацию о способе погрузки и упаковке для каждого товара
                    # Используем комбинацию как ключ, чтобы избежать дублирования
                    loading_key = f"{loading_method_value}|{packaging_kg_value}"
                    if loading_key not in [f"{p.get('loading_method', '')}|{p.get('packaging_kg', '')}" for p in products_loading]:
                        products_loading.append({
                            "loading_method": loading_method_value,
                            "packaging_kg": packaging_kg_value
                        })
    
    # Преобразуем словарь в список
    products_list = [{"name": name, "kg": str(int(kg)) if kg == int(kg) else str(kg)} for name, kg in products_dict.items()]
    
    # Формируем "Адрес отгрузки спецификации" / "Адрес погрузки"
    if all_fields:
        try:
            shipping_spec_address_text = build_shipping_spec_address_text(all_fields, products_list)
            if shipping_spec_address_text:
                table_fields_map["Адрес отгрузки спецификации"] = shipping_spec_address_text
                table_fields_map["Адрес отгрузки спецификации".strip()] = shipping_spec_address_text
        except Exception as e:
            logger.error(f"Error building shipping spec address text: {e}", exc_info=True)
    
    # Формируем описание способа отгрузки для всех товаров (для ${Товар отгружается})
    # Если галочка id: 131 ("Товар отгружается (в чем?)") установлена, текст будет взят из id 135 в process_word_template
    # Если галочка НЕ установлена, используем автоматическое формирование
    loading_description = get_loading_description_multiple(products_loading)
    if loading_description:
        table_fields_map["Товар отгружается"] = loading_description
        table_fields_map["Товар отгружается".strip()] = loading_description
    
    # Формируем "В соответствии с наименованием" только если галочка для ручного текста НЕ установлена
    # Если галочка установлена, текст будет взят из поля id 130 в process_word_template
    if not use_manual_product_description:
        product_description = get_product_description(products_list)
        if product_description:
            # Убеждаемся, что в конце есть точка
            if not product_description.endswith('.'):
                product_description = product_description.rstrip('.') + "."
            table_fields_map["В соответствии с наименованием"] = product_description
            table_fields_map["В соответствии с наименованием".strip()] = product_description
    
    # Формируем "Дата отгрузки" из таблицы для использования, если [SYSTEM] DATE LOADING не заполнен
    shipping_date_text = generate_shipping_date_text(table_fields)
    logger.info(f"Generated shipping date text from table: {repr(shipping_date_text)}")
    
    # Добавляем "Дата отгрузки" из таблицы
    table_fields_map["Дата отгрузки"] = shipping_date_text
    table_fields_map["Дата отгрузки".strip()] = shipping_date_text
    
    # Добавляем [SYSTEM] DATE LOADING из таблицы (будет использоваться, если не заполнен в верхнем уровне)
    table_fields_map["[SYSTEM] DATE LOADING"] = shipping_date_text
    table_fields_map["[SYSTEM] DATE LOADING".strip()] = shipping_date_text
    
    return table_fields_map


def find_field_by_id(fields_list: list[dict[str, Any]], target_id: int) -> dict[str, Any] | None:
    """Рекурсивно ищет поле по id"""
    for field in fields_list:
        field_id = field.get("id")
        if field_id == target_id:
            return field
        
        # Проверяем вложенные поля в title
        if field.get("type") == "title":
            field_value = field.get("value")
            if isinstance(field_value, dict):
                nested_fields = field_value.get("fields", [])
                if nested_fields:
                    found = find_field_by_id(nested_fields, target_id)
                    if found:
                        return found
        
        # Проверяем табличные значения (type="table")
        if field.get("type") == "table":
            table_value = field.get("value", [])
            if isinstance(table_value, list):
                for row in table_value:
                    if not isinstance(row, dict):
                        continue
                    cells = row.get("cells", [])
                    if not isinstance(cells, list):
                        continue
                    for cell in cells:
                        if not isinstance(cell, dict):
                            continue
                        if cell.get("id") == target_id:
                            return cell
    return None


def find_field_by_name(
    fields_list: list[dict[str, Any]], target_name: str
) -> dict[str, Any] | None:
    """Рекурсивно ищет поле/ячейку по точному имени (включая table cells)."""
    for field in fields_list:
        name = field.get("name")
        if name == target_name:
            return field

        # title nesting
        if field.get("type") == "title":
            field_value = field.get("value")
            if isinstance(field_value, dict):
                nested_fields = field_value.get("fields", [])
                if nested_fields:
                    found = find_field_by_name(nested_fields, target_name)
                    if found:
                        return found

        # table cells nesting
        if field.get("type") == "table":
            table_value = field.get("value", [])
            if isinstance(table_value, list):
                for row in table_value:
                    if not isinstance(row, dict):
                        continue
                    cells = row.get("cells", [])
                    if not isinstance(cells, list):
                        continue
                    for cell in cells:
                        if not isinstance(cell, dict):
                            continue
                        if cell.get("name") == target_name:
                            return cell

    return None


def check_manual_shipping_spec_address(fields: list[dict[str, Any]]) -> tuple[bool, str]:
    """
    Проверяет, установлена ли галочка "Адрес отгрузки спецификации (Текст вручную)" (id: 145)
    и возвращает текст из поля "Адрес отгрузки спецификации (текст)" (id: 146).
    Используется для плейсхолдера ${Адрес отгрузки спецификации}.
    
    Returns:
        tuple[bool, str]: (is_manual, manual_text) - установлена ли галочка и текст из поля id 146
    """
    field_145 = find_field_by_id(fields, 145)
    is_manual = False
    if field_145:
        field_value = field_145.get("value")
        if isinstance(field_value, str):
            is_manual = field_value.lower() == "checked"
        elif isinstance(field_value, bool):
            is_manual = field_value
        elif isinstance(field_value, int):
            is_manual = field_value > 0
    
    manual_text = ""
    if is_manual:
        field_146 = find_field_by_id(fields, 146)
        if field_146:
            manual_text = extract_field_value(field_146)
            if manual_text:
                manual_text = str(manual_text).strip()
    
    return is_manual, manual_text


def build_shipping_spec_address_text(
    all_fields: list[dict[str, Any]],
    products_list: list[dict[str, Any]] | None = None,
) -> str:
    """
    Формирует текст для плейсхолдера ${Адрес отгрузки спецификации} с учетом:
    - галочки "Адрес отгрузки спецификации (Текст вручную)" (id: 145) и поля id: 146
    - значения поля id: 133 (multiple_choice) с возможностью выбора нескольких адресов
    - списка товаров products_list (опционально) для формирования вида "Товар: Адрес"
    """
    # 1. Проверяем ручной режим
    is_manual, manual_text = check_manual_shipping_spec_address(all_fields)
    if is_manual and manual_text:
        return manual_text
    
    # 2. Получаем значения поля id 133
    field_133 = find_field_by_id(all_fields, 133)
    if not field_133:
        return ""
    
    value = field_133.get("value")
    addresses_short: list[str] = []
    
    if isinstance(value, dict) and "choice_names" in value:
        choice_names = value.get("choice_names", [])
        for name in choice_names:
            if name:
                addresses_short.append(str(name).strip())
    elif value:
        addresses_short.append(str(value).strip())
    
    if not addresses_short:
        return ""
    
    # 3. Преобразуем короткие названия в полные адреса
    full_addresses: list[str] = []
    for short in addresses_short:
        try:
            full = get_full_address(short)
        except Exception:
            full = short
        if full:
            full_addresses.append(str(full).strip())
    
    if not full_addresses:
        return ""
    
    # 4. Если нет товаров или адрес один — ведем себя как раньше
    if not products_list or len(full_addresses) == 1:
        if len(full_addresses) == 1:
            return full_addresses[0]
        # несколько адресов, но нет товаров — просто перечисляем
        return "; ".join(full_addresses)
    
    # 5. Несколько адресов и есть список товаров:
    #    формируем строки вида "Товар: Адрес"
    lines: list[str] = []
    for idx, product in enumerate(products_list):
        product_name = str(product.get("name", "")).strip()
        if not product_name:
            continue
        if idx < len(full_addresses):
            addr = full_addresses[idx]
        else:
            # если адресов меньше, чем товаров, используем последний адрес
            addr = full_addresses[-1]
        lines.append(f"{product_name}: {addr}")
    
    return "\n".join(lines)


def check_manual_loading_in_what(fields: list[dict[str, Any]]) -> tuple[bool, str]:
    """
    Проверяет, установлена ли галочка "Товар отгружается (в чем?)" (id: 134)
    и возвращает текст из поля "Товар отгружается (в чем?) текст" (id: 135)
    Используется для плейсхолдера ${Товар отгружается}
    
    Returns:
        tuple[bool, str]: (is_manual, manual_text) - установлена ли галочка и текст из поля id 135
    """
    # Ищем поле id 134 (checkmark)
    field_134 = find_field_by_id(fields, 134)
    is_manual = False
    if field_134:
        field_value = field_134.get("value")
        # Проверяем, установлена ли галочка
        if isinstance(field_value, str):
            is_manual = field_value.lower() == "checked"
        elif isinstance(field_value, bool):
            is_manual = field_value
        elif isinstance(field_value, int):
            is_manual = field_value > 0
    
    # Если галочка установлена, ищем поле id 135
    manual_text = ""
    if is_manual:
        field_135 = find_field_by_id(fields, 135)
        if field_135:
            manual_text = extract_field_value(field_135)
            if manual_text:
                manual_text = str(manual_text).strip()
    
    return is_manual, manual_text


def check_manual_loading_text(fields: list[dict[str, Any]]) -> tuple[bool, str]:
    """
    Проверяет, установлена ли галочка для ручного текста "В соответствии с наименованием"
    и возвращает текст из поля "Товар отгружается в соответствии с наименованием товара" (id: 130)
    Используется для плейсхолдера ${В соответствии с наименованием}
    
    Returns:
        tuple[bool, str]: (is_manual, manual_text) - установлена ли галочка и текст из поля id 130
    """
    # Ищем поле id 131 (checkmark) - но это может быть другая галочка для "В соответствии с наименованием"
    # Пока используем ту же логику, что и раньше
    field_131 = find_field_by_id(fields, 131)
    is_manual = False
    if field_131:
        field_name = field_131.get("name", "")
        # Проверяем, что это галочка для "В соответствии с наименованием"
        if "Текст вручную" in field_name or "в соответствии" in field_name.lower():
            field_value = field_131.get("value")
            # Проверяем, установлена ли галочка
            if isinstance(field_value, str):
                is_manual = field_value.lower() == "checked"
            elif isinstance(field_value, bool):
                is_manual = field_value
            elif isinstance(field_value, int):
                is_manual = field_value > 0
    
    # Если галочка установлена, ищем поле id 130
    manual_text = ""
    if is_manual:
        field_130 = find_field_by_id(fields, 130)
        if field_130:
            manual_text = extract_field_value(field_130)
            if manual_text:
                manual_text = str(manual_text).strip()
    
    return is_manual, manual_text


def process_word_template(
    template_path: str | Path,
    output_path: str | Path,
    fields: list[dict[str, Any]],
    director_fio: str = "ВЫ НЕ УКАЗАЛИ ДИРЕКТОРА",
    is_general_director: bool = False
):
    """Обрабатывает шаблон Word и создает новый файл с подставленными значениями"""
    doc = Document(template_path)
    
    fields_map = create_fields_map(fields)
    
    # --- Подписи: покупатель и поставщик ---
    # ${FinalStringSupplier} формируется по ФИО Поставщика (id: 117).
    supplier_field_117 = find_field_by_id(fields, 117)
    supplier_value = (
        extract_field_value(supplier_field_117)
        if supplier_field_117 is not None
        else ""
    )
    if not supplier_value:
        # fallback: иногда поле уже попадает в fields_map по имени
        supplier_value = fields_map.get("ФИО Поставщика") or ""

    supplier_value_str = str(supplier_value).strip()
    supplier_upper = supplier_value_str.upper()
    # Если это не ИП и в начале нет "ООО", добавляем форму "ООО «... »"
    if "ИП" not in supplier_upper and not supplier_upper.startswith("ООО"):
        supplier_value_str = f"ООО «{supplier_value_str}»"

    # ${FinalStringDirector} формируется по полю "Организация" (покупатель).
    # Поле может лежать внутри table cells, поэтому ищем рекурсивно.
    buyer_field = find_field_by_name(fields, "Организация")
    buyer_value = extract_field_value(buyer_field) if buyer_field is not None else ""
    if not buyer_value:
        # fallback на старые данные директора (на случай если в payload еще старый формат)
        buyer_value = fields_map.get("Организация") or director_fio

    # Формируем строки подписей.
    # Для поставщика (${FinalStringSupplier}) ячейка уже, поэтому уменьшаем max_length
    supplier_string = format_director_string(
        supplier_value_str,
        False,
        max_length=38,
    )
    buyer_string = format_director_string(str(buyer_value).strip(), False, max_length=42)

    # Добавляем данные директора/подписанта для уже существующих плейсхолдеров
    fields_map["[SYSTEM] ФИО ДИРЕКТОРА ОРГАНИЗАЦИИ"] = str(buyer_value).strip()
    fields_map["[SYSTEM] ФИО ДИРЕКТОРА ОРГАНИЗАЦИИ".strip()] = str(buyer_value).strip()

    fields_map["Director String"] = buyer_string
    fields_map["Director String".strip()] = buyer_string

    # Плейсхолдер `Director` (отдельное слово "Директор" в шаблоне)
    buyer_has_ip = "ИП" in str(buyer_value).upper()
    fields_map["Director"] = "" if buyer_has_ip else "Директор"

    # Новые плейсхолдеры
    fields_map["FinalStringDirector"] = buyer_string
    fields_map["FinalStringDirector".strip()] = buyer_string
    fields_map["FinalStringSupplier"] = supplier_string
    fields_map["FinalStringSupplier".strip()] = supplier_string
    table_fields = find_table_fields(fields)
    
    table_fields_map = extract_table_fields_to_map(table_fields, all_fields=fields)
    
    # Обрабатываем ${Товар отгружается}
    # Проверяем, установлена ли галочка "Товар отгружается (в чем?)" (id: 131)
    # Если да, используем текст из поля id 135 вместо автоматически сформированного
    is_manual_in_what, manual_in_what_text = check_manual_loading_in_what(fields)
    if is_manual_in_what and manual_in_what_text:
        logger.info(f"Using manual 'Товар отгружается' text from field id 135: {manual_in_what_text}")
        table_fields_map["Товар отгружается"] = manual_in_what_text
        table_fields_map["Товар отгружается".strip()] = manual_in_what_text
    
    # Обрабатываем ${В соответствии с наименованием}
    # Проверяем, установлена ли галочка для ручного текста "В соответствии с наименованием"
    # Если да, используем текст из поля id 130 вместо автоматически сформированного
    is_manual_product, manual_product_text = check_manual_loading_text(fields)
    if is_manual_product and manual_product_text:
        logger.info(f"Using manual 'В соответствии с наименованием' text from field id 130: {manual_product_text}")
        table_fields_map["В соответствии с наименованием"] = manual_product_text
        table_fields_map["В соответствии с наименованием".strip()] = manual_product_text
        # Добавляем также для полного названия поля (на случай если в шаблоне используется полное название)
        table_fields_map["Товар отгружается в соответствии с наименованием товара"] = manual_product_text
        table_fields_map["Товар отгружается в соответствии с наименованием товара".strip()] = manual_product_text
    
    # Если [SYSTEM] DATE LOADING не заполнен в верхнем уровне, используем значение из таблицы
    if "[SYSTEM] DATE LOADING" not in fields_map or not fields_map.get("[SYSTEM] DATE LOADING"):
        if "[SYSTEM] DATE LOADING" in table_fields_map and table_fields_map["[SYSTEM] DATE LOADING"]:
            fields_map["[SYSTEM] DATE LOADING"] = table_fields_map["[SYSTEM] DATE LOADING"]
            fields_map["[SYSTEM] DATE LOADING".strip()] = table_fields_map["[SYSTEM] DATE LOADING"]
            logger.info(f"Using [SYSTEM] DATE LOADING from table: {repr(table_fields_map['[SYSTEM] DATE LOADING'])}")
        else:
            logger.info("[SYSTEM] DATE LOADING not found in table_fields_map or is empty")
    else:
        logger.info(f"Using [SYSTEM] DATE LOADING from top-level fields: {repr(fields_map.get('[SYSTEM] DATE LOADING'))}")
    
    # Обновляем fields_map остальными значениями из таблицы (но не перезаписываем [SYSTEM] DATE LOADING)
    table_fields_map_for_update = {k: v for k, v in table_fields_map.items() if k != "[SYSTEM] DATE LOADING" and k != "[SYSTEM] DATE LOADING".strip()}
    fields_map.update(table_fields_map_for_update)
    
    logger.info("Processing paragraphs")
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        should_remove = replace_placeholders_in_paragraph(paragraph, fields_map)
        if should_remove:
            paragraphs_to_remove.append(paragraph)
    
    logger.info(f"Removing {len(paragraphs_to_remove)} empty paragraphs")
    for paragraph in paragraphs_to_remove:
        if paragraph._element is not None and paragraph._element.getparent() is not None:
            paragraph._element.getparent().remove(paragraph._element)
    
    logger.info("Starting table data insertion")
    
    for table_field in table_fields:
        table_name = table_field.get("name", "")
        table_value = table_field.get("value", [])
        
        if isinstance(table_value, list) and table_value:
            logger.info(f"Processing table '{table_name}' with {len(table_value)} rows")
            try:
                insert_table_data(doc, table_name, table_value)
                logger.info(f"Successfully processed table '{table_name}'")
            except Exception as e:
                logger.error(f"Error processing table '{table_name}': {e}", exc_info=True)
                raise
    
    logger.info("Processing remaining table cells with placeholders")
    for table in doc.tables:
        for row in table.rows:
            is_total_row = False
            has_placeholders = False
            
            # Проверяем, есть ли плейсхолдеры в строке
            for cell in row.cells:
                cell_text = cell.text
                if "Итого" in cell_text and "${" in cell_text:
                    is_total_row = True
                if "${" in cell_text:
                    has_placeholders = True
            
            # Обрабатываем только строки с плейсхолдерами, которые еще не обработаны
            # (строки, обработанные в insert_table_data, уже не содержат плейсхолдеров)
            if not is_total_row and has_placeholders:
                for cell in row.cells:
                    replace_placeholders_in_table_cell(cell, fields_map)
    
    logger.info(f"Saving document to {output_path}")
    doc.save(output_path)
    logger.info("Document saved successfully")
